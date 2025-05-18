import PyPDF2
import docx
import io
import tempfile
import os
import streamlit as st

def extract_text_from_pdf(pdf_file):
    """Enhanced PDF text extraction with error handling and text cleanup."""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:  # Only add non-empty pages
                text += page_text + "\n\n"
        
        # Basic text cleanup
        text = text.replace('\n\n', '\n').strip()
        return text
    except Exception as e:
        st.error(f"PDF extraction error: {str(e)}")
        return ""

def extract_text_from_docx(docx_file):
    """Enhanced DOCX extraction including tables and formatting."""
    try:
        doc = docx.Document(docx_file)
        full_text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    except Exception as e:
        st.error(f"DOCX extraction error: {str(e)}")
        return ""